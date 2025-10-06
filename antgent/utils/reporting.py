import io
from typing import Any

from docx import Document
from docx.document import Document as DocumentClass
from htmldocx import HtmlToDocx  # type: ignore
from markdown_it import MarkdownIt


def _add_metadata_sections(document: DocumentClass, metadata_sections: dict[str, str]) -> None:
    """Add metadata sections to document."""
    for title, content in metadata_sections.items():
        document.add_heading(title, level=2)
        document.add_paragraph(content)


def _make_header_cells_bold(cells) -> None:
    """Make all runs in header cells bold."""
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def _add_table_to_document(document: DocumentClass, title: str, table_data: list[dict[str, Any]]) -> None:
    """Add a single table to the document."""
    if not table_data:
        return

    document.add_heading(title, level=2)
    headers = list(table_data[0].keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Set headers and make them bold
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    _make_header_cells_bold(hdr_cells)

    # Add data rows
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row_data.get(header, ""))


def _add_tables(document: DocumentClass, tables: dict[str, list[dict[str, Any]]]) -> None:
    """Add all tables to document."""
    for title, table_data in tables.items():
        _add_table_to_document(document, title, table_data)


def markdown_to_docx_bytes(
    main_content_md: str,
    metadata_sections: dict[str, str] | None = None,
    tables: dict[str, list[dict[str, Any]]] | None = None,
    main_content_title: str = "Content",
    metadata_title: str = "Details",
) -> bytes:
    """
    Converts markdown text, metadata, and tables into a DOCX file as bytes.

    Args:
        main_content_md: The main content in Markdown format.
        metadata_sections: A dictionary where keys are titles for metadata sections
                           and values are their string content.
        tables: A dictionary where keys are titles for tables and values are a
                list of dictionaries representing rows.
        main_content_title: The main title for the markdown content section.
        metadata_title: The main title for the metadata/tables page.

    Returns:
        The generated DOCX file as bytes.
    """
    document = Document()

    if metadata_sections or tables:
        document.add_heading(metadata_title, level=1)

        if metadata_sections:
            _add_metadata_sections(document, metadata_sections)

        if tables:
            _add_tables(document, tables)

        document.add_page_break()

    document.add_heading(main_content_title, level=1)

    md = MarkdownIt()
    html_content = md.render(main_content_md)

    parser = HtmlToDocx()
    parser.add_html_to_document(html_content, document)

    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()
